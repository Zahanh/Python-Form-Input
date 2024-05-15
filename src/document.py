#from docx import Document
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class new_document:
    def __init__(self,doc_path:os.path,doc_name:str='example.docx') -> None:
        self.doc        = docx.Document()
        self.doc_path   = doc_path
        self.doc_name   = doc_name
        self.font       = 'Arial'

        title                   = self.doc.add_heading('D.W. Hirst & Associates INC', level=1)
        title.alignment         = WD_ALIGN_PARAGRAPH.CENTER
        title.style.font.name   = self.font
        title.style.font.size   = Pt(30)

    def input(self,b_email,b_phone,b_address,f_num,estimate,request,cust,c_phone,c_email,request_date,deliver,closing_date,parcel,p_address,subdivision,legal,buyer,lender,title,title_ins,flood_zone,notes):
        self.file_information(file_number=f_num,estimate=estimate)
        self.buissness_info(email=b_email,phone_number=b_phone,mailing_address=b_address)
        self.contact(request=request,customer=cust,phone=c_phone,email=c_email,request_date=request_date,deliver=deliver,closing=closing_date)
        self.property(parcel=parcel,address=p_address,subdivision=subdivision,legal=legal)
        self.certs(buyer=buyer,lender=lender,title=title,title_ins=title_ins,flood_zone=flood_zone,notes=notes)
        self.save()

    def buissness_info(self,email:str='',phone_number:str='',mailing_address:str=''):
        self.doc.add_heading(f'Survey Contact Information',level=3)

        p               = self.doc.add_paragraph()
        p.add_run(f'Phone:\t {phone_number}')
        p.add_run(f'\n')
        p.add_run(f'Email:\t {email}')
        p.add_run(f'\n')
        p.add_run(f'Mailing Address: ')
        p.add_run(f'{mailing_address}').bold=True
        self.doc.add_paragraph("_________________________________________________________________________________________________________")

    def file_information(self,file_number:str='',estimate:str=''):
        self.doc.add_paragraph("_________________________________________________________________________________________________________")
        heading2 = self.doc.add_heading(f'File Information',level=3)
        p = self.doc.add_paragraph()
        p.add_run(f'File: {file_number}\n')
        p.add_run(f'Estimate: ${estimate}')

    def contact(self,request:str='',customer:str='',phone:str='',email:str='',request_date:str='',deliver:str='',closing:str=''):
        self.doc.add_heading(f'Contact & Order Information',level=3)

        p = self.doc.add_paragraph()
        p.add_run(f'Requested by: {request}\n')
        p.add_run(f'Representing/Customer: {customer}\n')
        p.add_run(f'Phone Number: {phone}\n')
        p.add_run(f'Email: {email}\n')
        p.add_run(f'Date Requested: {request_date}\n')
        p.add_run(f'Deliver by: {deliver}\n')
        p.add_run(f'Closing Date: {closing}')

    def property(self,parcel:str='',address:str='',subdivision:str='',legal:str=''):
        self.doc.add_heading(f'Surveyed Property Information',level=3)

        p = self.doc.add_paragraph()
        p.add_run(f'Parcel Number: {parcel}\n')
        p.add_run(f'Parcel Address: {address}\n')
        p.add_run(f'Subdivision: {subdivision}\n')
        p.add_run(f'Legal Information: {legal}')

    def certs(self,buyer:str='',lender:str='',title:str='',title_ins:str='',flood_zone:str='',notes:str=''):
        self.doc.add_heading(f'Certification Information',level=3)

        p = self.doc.add_paragraph()
        p.add_run(f'Buyer: {buyer}\n')
        p.add_run(f'Lender: {lender}\n')
        p.add_run(f'Title: {title}\n')
        p.add_run(f'Title Ins Comp: {title_ins}\n')
        p.add_run(f'Flood Zone Information: {flood_zone}\n\n')
        p.add_run(f'General Notes: {notes}')
    
    def save(self):
        os.chdir(self.doc_path)
        self.doc.save(self.doc_name)
